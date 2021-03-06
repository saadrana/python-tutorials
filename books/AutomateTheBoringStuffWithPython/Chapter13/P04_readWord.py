# This program uses the python-docx module to manipulate Word documents

import docx

# Reading Word Documents
doc = docx.Document("demo.docx")
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)

print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)
print(doc.paragraphs[1].runs[4].text)

# Getting the Full Text from a .docx File
import P05_readDocx  # Don't do this, imports should be at the top of the file
print(P05_readDocx.getText("demo.docx"))
